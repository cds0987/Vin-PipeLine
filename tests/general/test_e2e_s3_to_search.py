"""
End-to-end integration test: MinIO (S3) → S3Scanner → Pipeline → Qdrant → RetrievalService

Simulates the full production flow for every supported file format.

Stages
──────
  0. Upload test files to MinIO bucket (simulates S3)
  1. S3Scanner discovers files → IngestJob[]
  2. pipeline.run() ingests each job  (parse → clean → chunk → embed → index)
  3. Verify all documents indexed in metadata store
  4. RetrievalService.search() returns results from Qdrant

Run:
    docker compose run --rm test pytest tests/general/test_e2e_s3_to_search.py -v -m "minio and qdrant"
"""
from __future__ import annotations

import io
import os
import uuid
from pathlib import Path

import pytest

import config
from utils.ai_provider import MockAIProvider
from utils.stores import InMemoryMetadataStore

pytestmark = [pytest.mark.minio, pytest.mark.qdrant, pytest.mark.e2e]

# ── constants ────────────────────────────────────────────────────────────────

_BUCKET = os.getenv("MINIO_TEST_BUCKET", "rag-pipeline-local")
_MINIO_URL = os.getenv("MINIO_TEST_ENDPOINT", "http://minio:9000")
_MINIO_USER = os.getenv("MINIO_ROOT_USER", "minioadmin")
_MINIO_PASS = os.getenv("MINIO_ROOT_PASSWORD", "minioadmin")


# ── test file factories ───────────────────────────────────────────────────────


def _txt_bytes() -> bytes:
    return (
        "Chính sách nghỉ phép VSF 2024\n\n"
        "Nhân viên chính thức được hưởng 15 ngày phép năm có lương.\n"
        "Phép năm không dùng hết không được chuyển sang năm sau.\n"
        "Đơn xin nghỉ phép phải nộp trước ít nhất 3 ngày làm việc.\n"
        "Trường hợp nghỉ đột xuất phải thông báo qua email cho quản lý trực tiếp.\n"
    ).encode("utf-8")


def _md_bytes() -> bytes:
    return (
        "# Sổ tay nhân viên VSF\n\n"
        "## Phúc lợi\n\n"
        "- Bảo hiểm sức khoẻ toàn diện cho nhân viên và gia đình\n"
        "- Thưởng cuối năm theo hiệu suất làm việc\n"
        "- Hỗ trợ chi phí đào tạo và phát triển nghề nghiệp\n\n"
        "## Giờ làm việc\n\n"
        "Giờ làm việc tiêu chuẩn: 8h00 đến 17h30, thứ Hai đến thứ Sáu.\n"
        "Làm thêm giờ được tính phụ cấp theo quy định Bộ Lao động.\n"
    ).encode("utf-8")


def _html_bytes() -> bytes:
    return (
        "<html><body>"
        "<h1>Chính sách công tác</h1>"
        "<p>Nhân viên được hoàn tiền vé máy bay tối đa <b>5.000.000 VNĐ</b> mỗi chuyến.</p>"
        "<p>Khách sạn: mức tối đa 800.000 VNĐ/đêm tại thành phố lớn.</p>"
        "<p>Chứng từ hoàn tiền nộp trong vòng 7 ngày sau khi công tác về.</p>"
        "</body></html>"
    ).encode("utf-8")


def _docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Quy định bảo mật thông tin VSF", 0)
    doc.add_heading("1. Phạm vi áp dụng", 1)
    doc.add_paragraph(
        "Quy định này áp dụng cho toàn bộ nhân viên, cộng tác viên và đối tác "
        "có quyền truy cập vào hệ thống thông tin của Tập đoàn VSF."
    )
    doc.add_heading("2. Nghĩa vụ của nhân viên", 1)
    doc.add_paragraph(
        "Không chia sẻ thông tin nội bộ, tài liệu mật với bên ngoài dưới bất kỳ hình thức nào. "
        "Vi phạm có thể dẫn đến chấm dứt hợp đồng lao động và truy cứu trách nhiệm pháp lý."
    )
    doc.add_heading("3. Mật khẩu và tài khoản", 1)
    doc.add_paragraph(
        "Mật khẩu phải có ít nhất 12 ký tự, bao gồm chữ hoa, chữ thường, số và ký tự đặc biệt. "
        "Không được chia sẻ mật khẩu với bất kỳ ai, kể cả đồng nghiệp hay quản lý."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_bytes() -> bytes:
    """Image-based PDF (no text layer) — pipeline uses OCR fallback."""
    from PIL import Image, ImageDraw

    try:
        from PIL import ImageFont
        font = ImageFont.load_default(size=28)
    except TypeError:
        font = ImageFont.load_default()

    img = Image.new("RGB", (1000, 400), "white")
    draw = ImageDraw.Draw(img)
    draw.text((50, 60), "Báo cáo tài chính Q1/2024 — VSF Corporation", fill="black", font=font)
    draw.text((50, 130), "Doanh thu: 2.450 tỷ đồng  (+18% so với cùng kỳ)", fill="black", font=font)
    draw.text((50, 200), "Lợi nhuận sau thuế: 310 tỷ đồng", fill="black", font=font)

    buf = io.BytesIO()
    img.save(buf, format="PDF", resolution=120)
    return buf.getvalue()


def _png_bytes() -> bytes:
    """PNG image — pipeline uses direct OCR."""
    from PIL import Image, ImageDraw

    try:
        from PIL import ImageFont
        font = ImageFont.load_default(size=24)
    except TypeError:
        font = ImageFont.load_default()

    img = Image.new("RGB", (800, 200), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 40), "Sơ đồ tổ chức VSF: CEO → CTO, CFO, CHRO", fill="black", font=font)
    draw.text((40, 110), "Tổng nhân sự: 6.000 nhân viên tại 12 công ty con", fill="black", font=font)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _jpg_bytes() -> bytes:
    """JPEG image — pipeline uses direct OCR."""
    from PIL import Image, ImageDraw

    try:
        from PIL import ImageFont
        font = ImageFont.load_default(size=24)
    except TypeError:
        font = ImageFont.load_default()

    img = Image.new("RGB", (800, 200), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 60), "Biểu đồ KPI Q1 2024: đạt 94% mục tiêu doanh thu", fill="black", font=font)

    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=90)
    return buf.getvalue()


# ── shared file map ────────────────────────────────────────────────────────────

# key = S3 path relative to prefix (includes document_type sub-folder)
# value = factory function
_FILE_FACTORIES: dict[str, callable] = {
    "hr/leave_policy.txt":           _txt_bytes,
    "docs/employee_handbook.md":     _md_bytes,
    "web/travel_policy.html":        _html_bytes,
    "legal/security_policy.docx":    _docx_bytes,
    "reports/q1_financial.pdf":      _pdf_bytes,
    "diagrams/org_chart.png":        _png_bytes,
    "charts/kpi_q1.jpg":             _jpg_bytes,
}


# ── fixtures ──────────────────────────────────────────────────────────────────


@pytest.fixture(scope="module")
def minio_client():
    import boto3
    client = boto3.client(
        "s3",
        endpoint_url=_MINIO_URL,
        aws_access_key_id=_MINIO_USER,
        aws_secret_access_key=_MINIO_PASS,
    )
    existing = {b["Name"] for b in client.list_buckets().get("Buckets", [])}
    if _BUCKET not in existing:
        client.create_bucket(Bucket=_BUCKET)
    return client


# ── main E2E test ─────────────────────────────────────────────────────────────


def test_e2e_s3_scan_to_search(minio_client, monkeypatch):
    """
    Full flow: MinIO upload → S3Scanner → pipeline ingestion → Qdrant → search.
    Uses MockAIProvider (no external AI calls, fast, deterministic).
    """
    prefix = f"e2e-{uuid.uuid4().hex}/"

    # ── Stage 0: upload test files to MinIO ──────────────────────────────────
    uploaded: dict[str, str] = {}   # filename → s3_uri
    for relative_key, factory in _FILE_FACTORIES.items():
        full_key = prefix + relative_key
        minio_client.put_object(Bucket=_BUCKET, Key=full_key, Body=factory())
        uploaded[Path(relative_key).name] = f"s3://{_BUCKET}/{full_key}"

    # ── Patch settings so scanner + storage read from MinIO ──────────────────
    monkeypatch.setattr(config.settings, "S3_BUCKET", _BUCKET)
    monkeypatch.setattr(config.settings, "SCAN_PREFIX", prefix)
    monkeypatch.setattr(config.settings, "S3_ENDPOINT", _MINIO_URL)
    monkeypatch.setattr(config.settings, "AWS_ACCESS_KEY_ID", _MINIO_USER)
    monkeypatch.setattr(config.settings, "AWS_SECRET_ACCESS_KEY", _MINIO_PASS)
    monkeypatch.setattr(config.settings, "SEARCH_SCORE_THRESHOLD", 0.0)

    # ── Stage 1: S3 scanner ──────────────────────────────────────────────────
    from adapters.s3_adapter import S3Scanner

    metadata_store = InMemoryMetadataStore()
    jobs = S3Scanner(metadata_store).scan()

    scanned_filenames = {Path(j.file_uri).name for j in jobs}
    assert scanned_filenames == set(uploaded.keys()), (
        f"Scanner mismatch.\n"
        f"  Expected: {sorted(uploaded.keys())}\n"
        f"  Got:      {sorted(scanned_filenames)}"
    )

    # Verify S3Scanner sets typed fields correctly
    for job in jobs:
        assert job.file_name is not None, f"file_name missing on {job.file_uri}"
        assert job.document_type != "general" or "general" in job.file_uri, (
            f"Expected document_type from path, got 'general' for {job.file_uri}"
        )
        assert job.s3_last_modified is not None

    # ── Stage 2: pipeline ingestion ──────────────────────────────────────────
    from pipeline.run import run as pipeline_run
    from utils.stores import QdrantStore

    ai = MockAIProvider()
    vector_store = QdrantStore()
    ingestion_results: dict[str, dict] = {}

    for job in jobs:
        result = pipeline_run(
            job,
            ai_provider=ai,
            vector_store=vector_store,
            metadata_store=metadata_store,
        )
        ingestion_results[job.file_uri] = result

    # ── Stage 3: verify all files indexed ────────────────────────────────────
    for job in jobs:
        result = ingestion_results[job.file_uri]
        fname = Path(job.file_uri).name

        assert result["status"] == "indexed", (
            f"{fname}: expected status=indexed, got {result}"
        )
        assert result["section_count"] >= 1, (
            f"{fname}: expected at least 1 section, got {result['section_count']}"
        )

        doc = metadata_store.get_document(job.doc_id)
        assert doc is not None, f"{fname}: document not found in metadata store"
        assert doc.status == "indexed", f"{fname}: doc.status={doc.status}"
        assert doc.section_count == result["section_count"]
        assert doc.file_path == job.file_uri
        assert doc.file_name == job.file_name
        assert doc.markdown_s3_uri

    # ── Stage 4: retrieval search ─────────────────────────────────────────────
    from retrieval.service import RetrievalService

    service = RetrievalService(ai_provider=ai, vector_store=vector_store)

    # With MockAI (hash-based embeddings) + threshold=0.0, any query returns
    # results — we verify plumbing, not semantic relevance.
    results = service.search("chính sách nhân viên", top_k=10)
    assert len(results) > 0, "search() returned empty after indexing documents"

    ingested_doc_ids = {j.doc_id for j in jobs}
    result_doc_ids = {r["document_id"] for r in results}
    assert result_doc_ids & ingested_doc_ids, (
        f"No search results came from indexed documents.\n"
        f"  Indexed doc_ids: {ingested_doc_ids}\n"
        f"  Result doc_ids:  {result_doc_ids}"
    )

    # Verify response fields are complete
    for r in results:
        assert "section_id" in r and r["section_id"]
        assert "section_content" in r and r["section_content"]
        assert "source_s3_uri" in r and r["source_s3_uri"].startswith("s3://")
        assert "markdown_s3_uri" in r and r["markdown_s3_uri"]
        assert "document_id" in r
        assert "score" in r

    # ── Cleanup: remove test vectors from Qdrant ──────────────────────────────
    for job in jobs:
        try:
            vector_store.delete(job.doc_id)
        except Exception:
            pass


# ── per-format detail tests ───────────────────────────────────────────────────
# These run independently and verify format-specific pipeline behaviour.


@pytest.mark.parametrize("relative_key,expected_type", [
    ("hr/leave_policy.txt",        "hr"),
    ("docs/employee_handbook.md",  "docs"),
    ("web/travel_policy.html",     "web"),
    ("legal/security_policy.docx", "legal"),
    ("reports/q1_financial.pdf",   "reports"),
    ("diagrams/org_chart.png",     "diagrams"),
    ("charts/kpi_q1.jpg",          "charts"),
])
def test_per_format_ingestion(
    minio_client, monkeypatch, relative_key: str, expected_type: str
):
    """Each file format is individually scanned, ingested and verified."""
    prefix = f"fmt-{uuid.uuid4().hex}/"
    fname = Path(relative_key).name
    suffix = Path(relative_key).suffix
    full_key = prefix + relative_key

    minio_client.put_object(
        Bucket=_BUCKET,
        Key=full_key,
        Body=_FILE_FACTORIES[relative_key](),
    )

    monkeypatch.setattr(config.settings, "S3_BUCKET", _BUCKET)
    monkeypatch.setattr(config.settings, "SCAN_PREFIX", prefix)
    monkeypatch.setattr(config.settings, "S3_ENDPOINT", _MINIO_URL)
    monkeypatch.setattr(config.settings, "AWS_ACCESS_KEY_ID", _MINIO_USER)
    monkeypatch.setattr(config.settings, "AWS_SECRET_ACCESS_KEY", _MINIO_PASS)

    from adapters.s3_adapter import S3Scanner
    from pipeline.run import run as pipeline_run
    from utils.stores import QdrantStore

    metadata_store = InMemoryMetadataStore()
    jobs = S3Scanner(metadata_store).scan()

    assert len(jobs) == 1, f"Expected 1 job for {fname}, got {len(jobs)}"
    job = jobs[0]

    assert job.file_name == fname, f"file_name mismatch: {job.file_name}"
    assert job.document_type == expected_type, (
        f"document_type: expected '{expected_type}', got '{job.document_type}'"
    )

    ai = MockAIProvider()
    vector_store = QdrantStore()

    result = pipeline_run(job, ai_provider=ai, vector_store=vector_store, metadata_store=metadata_store)

    assert result["status"] == "indexed", f"{fname}: {result}"
    assert result["section_count"] >= 1, f"{fname}: no sections produced"

    doc = metadata_store.get_document(job.doc_id)
    assert doc.status == "indexed"
    assert doc.file_type == suffix.lstrip(".")

    # Cleanup
    try:
        vector_store.delete(job.doc_id)
    except Exception:
        pass
